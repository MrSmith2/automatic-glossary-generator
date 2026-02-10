import streamlit as st
import pandas as pd
from datetime import datetime

from file_utils import read_file_content
from ollama_utils import check_ollama_connection, generate_definitions_batch
from term_extractor import extract_terms, get_term_context, filter_terms
from prompts import DEFINITION_PROMPT
from config import DEFAULT_COVERAGE_PERCENT, DEFAULT_LLM_CHUNK_SIZE, DEFAULT_SPACY_BATCH_SIZE, DEFAULT_SEED

st.title("Генератор глоссария терминов")

connection_status, available_models = check_ollama_connection()
if connection_status:
    st.success("Соединение с Ollama установлено")
else:
    st.error("Не удается подключиться к Ollama. Убедитесь, что Ollama запущен.")
    st.stop()

with st.sidebar:
    st.header("Настройки")
    
    if available_models:
        default_model = "qwen3:4b"
        if default_model not in available_models:
            default_model = available_models[0]
        model_name = st.selectbox(
            "Модель Ollama",
            available_models,
            index=available_models.index(default_model) if default_model in available_models else 0
        )
    else:
        model_name = st.text_input("Модель Ollama", "qwen3:4b")
    
    context_size = st.selectbox("Контекст модели", [4096, 8192, 16384], index=1)
    seed = st.number_input("Seed LLM", min_value=0, max_value=99999, value=DEFAULT_SEED)
    max_terms = st.slider("Макс. кол-во терминов", 10, 100, 50)
    context_window = st.slider("Окно контекста (символов)", 100, 500, 300)
    
    st.divider()
    st.subheader("Метод извлечения")
    
    extraction_method = st.selectbox(
        "Метод",
        ["llm", "spacy"],
        index=0
    )
    
    coverage_percent = st.slider(
        "Покрытие документа (%)",
        min_value=10,
        max_value=100,
        value=DEFAULT_COVERAGE_PERCENT,
        step=5
    )
    
    if extraction_method == "llm":
        chunk_size = st.select_slider(
            "Размер чанка LLM",
            options=[4000, 8000, 12000, 16000],
            value=DEFAULT_LLM_CHUNK_SIZE
        )
        batch_size = None
    else:
        batch_size = st.select_slider(
            "Размер батча spaCy",
            options=[50000, 100000, 150000],
            value=DEFAULT_SPACY_BATCH_SIZE,
            format_func=lambda x: f"{x//1000}k"
        )
        chunk_size = None

for key, default in [
    ('terms', []),
    ('definitions', []),
    ('document_text', ""),
    ('loaded_file_name', None),
    ('extraction_stats', None),
    ('terms_with_context', [])
]:
    if key not in st.session_state:
        st.session_state[key] = default

uploaded_file = st.file_uploader("Загрузите документ", type=['txt', 'docx', 'pdf'])

if uploaded_file:
    if st.session_state.loaded_file_name != uploaded_file.name:
        with st.spinner('Читаю документ...'):
            st.session_state.document_text = read_file_content(uploaded_file)
            st.session_state.loaded_file_name = uploaded_file.name
            st.session_state.terms = []
            st.session_state.definitions = []
            st.session_state.extraction_stats = None
            st.session_state.terms_with_context = []
        st.success(f"Загружен: {uploaded_file.name}")

if st.session_state.document_text:
    text = st.session_state.document_text
    text_len = len(text)
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Всего символов", f"{text_len:,}")
    with col2:
        chars_to_process = int(text_len * coverage_percent / 100)
        st.metric("Будет обработано", f"{chars_to_process:,}")
    
    with st.expander(f"Просмотр текста", expanded=False):
        st.text_area(
            "Полный текст документа",
            text,
            height=500,
            label_visibility="collapsed",
            disabled=True
        )

if st.button("Извлечь термины", disabled=not st.session_state.document_text):
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    def update_progress(current, total):
        progress = current / total if total > 0 else 0
        progress_bar.progress(progress)
        unit = "чанка" if extraction_method == "llm" else "батча"
        status_text.text(f"Обработка {unit} {current}/{total}...")
    
    with st.spinner(f'Извлекаю термины: {extraction_method.upper()}...'):
        terms, stats = extract_terms(
            st.session_state.document_text,
            top_n=max_terms * 2,
            method=extraction_method,
            model_name=model_name,
            context_size=context_size, 
            coverage_percent=coverage_percent,
            chunk_size=chunk_size,
            batch_size=batch_size,
            progress_callback=update_progress,
            seed=seed
        )
        st.session_state.terms = filter_terms(terms)[:max_terms]
        st.session_state.definitions = []
        st.session_state.extraction_stats = stats
        st.session_state.terms_with_context = []
    
    progress_bar.empty()
    status_text.empty()
    st.success(f"Найдено терминов: {len(st.session_state.terms)}")

if st.session_state.terms:
    st.subheader("Найденные термины")

    terms_txt = "\n".join(
        sorted(t["term"] for t in st.session_state.terms)
    )
    st.download_button(
        "Скачать список терминов (TXT)",
        terms_txt,
        f"terms_{datetime.now():%Y%m%d_%H%M%S}.txt",
        "text/plain"
    )
    
    sort_alpha = st.checkbox("Сортировать по алфавиту", value=True)
    
    for i in range(len(st.session_state.terms)):
        if f"term_{i}" not in st.session_state:
            st.session_state[f"term_{i}"] = True
    
    col_btn1, col_btn2, _ = st.columns([1, 1, 4])
    with col_btn1:
        if st.button("Выбрать все"):
            for i in range(len(st.session_state.terms)):
                st.session_state[f"term_{i}"] = True
            st.rerun()
    with col_btn2:
        if st.button("Снять все"):
            for i in range(len(st.session_state.terms)):
                st.session_state[f"term_{i}"] = False
            st.rerun()
    
    display_order = list(enumerate(st.session_state.terms))
    if sort_alpha:
        display_order.sort(key=lambda x: x[1]["term"].lower())
    
    selected_terms = []
    cols = st.columns(3)
    for j, (i, term_data) in enumerate(display_order):
        with cols[j % 3]:
            if st.checkbox(term_data["term"], key=f"term_{i}"):
                selected_terms.append(term_data)
    
    st.write(f"Выбрано: {len(selected_terms)} терминов")
    
    if st.button("Сгенерировать определения", type="primary", disabled=not selected_terms):
        terms_with_context = [
            {
                "term": t["term"],
                "context": get_term_context(st.session_state.document_text, t["term"], context_window)
            }
            for t in selected_terms
        ]
        
        progress = st.progress(0)
        status = st.empty()
        
        results = []
        for i, item in enumerate(terms_with_context):
            status.text(f"Генерация: {item['term']} ({i+1}/{len(terms_with_context)})")
            result = generate_definitions_batch([item], model_name, context_size, DEFINITION_PROMPT, seed=seed)
            result[0]["context"] = item["context"]
            results.extend(result)
            progress.progress((i + 1) / len(terms_with_context))
        
        st.session_state.definitions = results
        st.session_state.terms_with_context = terms_with_context
        progress.empty()
        status.text("Генерация завершена")

if st.session_state.definitions:
    st.subheader("Глоссарий")
    
    show_context = st.checkbox("Показывать контекст для LLM", value=False)
    
    edited_definitions = []
    
    for i, item in enumerate(st.session_state.definitions):
        term = item['term']
        definition = item.get('definition', '')
        context = item.get('context', '')
        error = item.get('error')
        
        st.markdown(f"### {term}")
        
        if show_context and context:
            with st.expander(f"Контекст для '{term}'", expanded=False):
                st.info(context)
        
        if error:
            st.warning(f"Ошибка: {error}")
        
        edited_def = st.text_area(
            f"Определение",
            value=definition,
            key=f"def_{i}",
            label_visibility="collapsed",
            height=100
        )
        edited_definitions.append({
            "term": term, 
            "definition": edited_def,
            "context": context
        })
        st.divider()
    
    st.subheader("Экспорт")
    
    col1, col2, col3, col4 = st.columns(4)
        
    with col1:
        df = pd.DataFrame([{"term": d["term"], "definition": d["definition"]} for d in edited_definitions])
        csv = df.to_csv(index=False, encoding='utf-8-sig')
        st.download_button(
            "CSV",
            csv,
            f"glossary_{datetime.now():%Y%m%d_%H%M%S}.csv",
            "text/csv"
        )
    
    with col2:
        md = "\n\n".join([
            f"**{d['term']}** — {d['definition']}"
            for d in edited_definitions if d['definition']
        ])
        st.download_button(
            "Markdown",
            md,
            f"glossary_{datetime.now():%Y%m%d_%H%M%S}.md",
            "text/markdown"
        )
    
    with col3:
        txt = "\n\n".join([
            f"{d['term']} — {d['definition']}"
            for d in edited_definitions if d['definition']
        ])
        st.download_button(
            "TXT",
            txt,
            f"glossary_{datetime.now():%Y%m%d_%H%M%S}.txt",
            "text/plain"
        )
    
    with col4:
        from io import BytesIO
        from docx import Document
        from docx.shared import Pt, RGBColor
        
        buf = BytesIO()
        doc = Document()
        doc.add_heading("Глоссарий", level=1)
        
        for d in edited_definitions:
            if not d['definition']:
                continue
            p = doc.add_paragraph()
            run_term = p.add_run(d['term'])
            run_term.bold = True
            run_term.font.size = Pt(12)
            p.add_run(f" — {d['definition']}")
        
        doc.save(buf)
        st.download_button(
            "DOCX",
            buf.getvalue(),
            f"glossary_{datetime.now():%Y%m%d_%H%M%S}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

with st.sidebar:
    st.divider()
    st.caption("**Текущие настройки:**")
    st.caption(f"Модель: {model_name}")
    st.caption(f"Контекст: {context_size}")
    st.caption(f"Seed: {seed}")
    st.caption(f"Метод: {extraction_method}")
    st.caption(f"Покрытие: {coverage_percent}%")
    if extraction_method == "llm":
        st.caption(f"Чанк: {chunk_size}")
    else:
        st.caption(f"Батч: {batch_size//1000}k")
    