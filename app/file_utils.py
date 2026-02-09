import io
from docx import Document
import pdfplumber

def docx_to_text(docx_file):
    doc = Document(io.BytesIO(docx_file.read()))
    full_text = []
    for p in doc.paragraphs:
        full_text.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    return '\n'.join(full_text)

def pdf_to_text(pdf_file):
    with pdfplumber.open(io.BytesIO(pdf_file.read())) as pdf:
        text = ""
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text

def read_file_content(file):
    file_type = file.type
    if file_type == "text/plain":
        return str(file.read(), "utf-8", errors="replace")
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return docx_to_text(file)
    elif file_type == "application/pdf":
        return pdf_to_text(file)
    else:
        file_name = file.name.lower()
        if file_name.endswith('.txt'):
            return str(file.read(), "utf-8", errors="replace")
        elif file_name.endswith('.docx'):
            return docx_to_text(file)
        elif file_name.endswith('.pdf'):
            return pdf_to_text(file)
        else:
            raise Exception(f"Неподдерживаемый формат файла: {file_type}")
